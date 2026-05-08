import sys
import subprocess

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()

    # Apply global styling: Times New Roman, Size 12, 1.5 line spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Modify Heading 1 style
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.line_spacing = 1.5

    # Modify Heading 2 style
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.line_spacing = 1.5

    # 1. Title Page
    for _ in range(5):
        doc.add_paragraph('')
        
    title = doc.add_paragraph('Project Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    doc.add_paragraph('\n\n')

    details = [
        "Course: CSE274 – Applied Machine Learning",
        "Names of Students: [Enter Names]",
        "Roll Numbers: [Enter Roll Numbers]",
        "Instructor Name: [Enter Instructor Name]",
        "Department / University: [Enter Dept/University]",
        "Submission Date: [Enter Date]"
    ]

    for d in details:
        p = doc.add_paragraph(d)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Helper for adding bullet points
    def add_bullets(items):
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    # 2. Abstract
    doc.add_heading('2. Abstract', level=1)
    doc.add_paragraph('Brief overview of the project (150–250 words).')
    add_bullets([
        'Problem statement',
        'Techniques used',
        'Key results'
    ])

    # 3. Introduction
    doc.add_heading('3. Introduction', level=1)
    add_bullets([
        'Background of the problem',
        'Importance of the study',
        'Real-world relevance',
        'Objective of the project'
    ])

    # 4. Problem Statement
    doc.add_heading('4. Problem Statement', level=1)
    doc.add_paragraph('Clearly define:')
    add_bullets([
        'What problem you are solving',
        'Type: Classification / Regression / Clustering'
    ])
    doc.add_paragraph('Example:\n- Disease prediction (Classification)\n- House price prediction (Regression)\n- Customer segmentation (Clustering)')

    # 5. Dataset Description
    doc.add_heading('5. Dataset Description', level=1)
    add_bullets([
        'Dataset source (Kaggle / UCI / etc.)',
        'Number of records and features',
        'Feature description (table format recommended)',
        'Target variable (if applicable)'
    ])

    # 6. Data Preprocessing
    doc.add_heading('6. Data Preprocessing', level=1)
    add_bullets([
        'Handling missing values',
        'Outlier detection & treatment',
        'Encoding categorical variables',
        'Feature scaling / normalization',
        'Handling class imbalance (if classification)',
        'Data leakage prevention'
    ])

    # 7. Feature Engineering & Dimensionality
    doc.add_heading('7. Feature Engineering & Dimensionality', level=1)
    doc.add_heading('Feature selection techniques:', level=2)
    add_bullets([
        'Variance Threshold',
        'Correlation-based removal',
        'Forward / Backward selection'
    ])
    add_bullets([
        'Feature extraction',
        'PCA / LDA (if used)',
        'Explanation of selected features'
    ])

    # 8. Methodology
    doc.add_heading('8. Methodology', level=1)
    doc.add_paragraph('(Choose based on project type)')
    
    doc.add_heading('A. For Classification', level=2)
    add_bullets(['Models used: Logistic Regression, Naïve Bayes, KNN / SVM / Decision Tree'])
    
    doc.add_heading('B. For Regression', level=2)
    add_bullets(['Models used: Linear Regression, Multiple Regression, Polynomial / Regularized models'])
    
    doc.add_heading('C. For Clustering', level=2)
    add_bullets(['Algorithms used: K-Means, Hierarchical Clustering, DBSCAN'])
    
    doc.add_paragraph('\nInclude:')
    add_bullets([
        'Reason for choosing each algorithm',
        'Workflow diagram (recommended)'
    ])

    # 9. Implementation Details
    doc.add_heading('9. Implementation Details', level=1)
    add_bullets([
        'Tools used: Python, Jupyter Notebook',
        'Libraries: NumPy, Pandas, Scikit-learn, Matplotlib, Seaborn',
        'Parameter settings: (e.g., K (clusters), Learning rate, Depth of tree)'
    ])

    # 10. Model Evaluation
    doc.add_heading('10. Model Evaluation', level=1)
    
    doc.add_heading('For Classification', level=2)
    add_bullets(['Confusion Matrix', 'Accuracy, Precision, Recall, F1-score', 'ROC Curve, AUC'])
    
    doc.add_heading('For Regression', level=2)
    add_bullets(['MAE, MSE, RMSE', 'R² Score', 'Residual plots'])
    
    doc.add_heading('For Clustering', level=2)
    add_bullets(['Silhouette Score', 'WCSS (Elbow Method)', 'Davies-Bouldin Index'])

    # 11. Results & Visualization
    doc.add_heading('11. Results & Visualization', level=1)
    add_bullets([
        'Graphs: ROC Curve, Elbow Graph, Cluster plots, Actual vs Predicted plots',
        'Tables comparing models'
    ])

    # 12. Hyperparameter Tuning
    doc.add_heading('12. Hyperparameter Tuning', level=1)
    add_bullets([
        'Grid Search / Random Search',
        'Cross-validation',
        'Best parameters found',
        'Performance improvement'
    ])

    # 13. Interpretation & Insights
    doc.add_heading('13. Interpretation & Insights', level=1)
    add_bullets([
        'What did the model learn?',
        'Key patterns or trends',
        'Business/real-world insights'
    ])

    # 14. Conclusion
    doc.add_heading('14. Conclusion', level=1)
    add_bullets([
        'Summary of findings',
        'Best performing model',
        'Limitations',
        'Future scope'
    ])

    # 15. Appendix
    doc.add_heading('15. Appendix', level=1)
    add_bullets([
        'Code snippets',
        'Additional graphs',
        'Screenshots'
    ])

    # 16. References
    doc.add_heading('16. References', level=1)
    add_bullets([
        'Dataset source',
        'Research papers / websites',
        'Books'
    ])

    # Save the document
    file_path = 'ML_Project_Report_Template.docx'
    doc.save(file_path)
    print(f"Success! Saved to {file_path}")

if __name__ == '__main__':
    main()
